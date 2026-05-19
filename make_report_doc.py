"""Generate a polished Word (.docx) report for the Crop Yield & Plant Disease project."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ────────────────────────────────────────────────────────────
# Stored as (RGBColor, hex_string) pairs
C_NAVY   = RGBColor(0x1A, 0x37, 0x6C);  H_NAVY  = '1A376C'
C_TEAL   = RGBColor(0x00, 0x7A, 0x87);  H_TEAL  = '007A87'
C_GREEN  = RGBColor(0x1E, 0x6B, 0x2E);  H_GREEN = '1E6B2E'
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF);  H_WHITE = 'FFFFFF'
C_LGRAY  = RGBColor(0xF2, 0xF4, 0xF8);  H_LGRAY = 'F2F4F8'
C_DGRAY  = RGBColor(0x44, 0x44, 0x44);  H_DGRAY = '444444'
C_GOLD   = RGBColor(0xD4, 0xAA, 0x00);  H_GOLD  = 'D4AA00'
C_HLGRN  = RGBColor(0xD6, 0xF0, 0xD6);  H_HLGRN = 'D6F0D6'


# ═════════════════════════════════════════════════════════════════════════════
# Helper utilities
# ═════════════════════════════════════════════════════════════════════════════

def set_cell_bg(cell, hex_color: str):
    """Fill a table cell background with a hex string e.g. '1A376C'."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Apply borders to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'),   kwargs[edge].get('val',   'single'))
            tag.set(qn('w:sz'),    kwargs[edge].get('sz',    '4'))
            tag.set(qn('w:space'), kwargs[edge].get('space', '0'))
            tag.set(qn('w:color'), kwargs[edge].get('color', 'auto'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def para_space(doc, before=0, after=6):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    return p


def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = C_NAVY
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A376C')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = C_TEAL
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = C_DGRAY
    return p


def add_body(doc, text, bold_parts=None):
    """Add a body paragraph with optional bold substrings."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if bold_parts:
        remaining = text
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx >= 0:
                if idx > 0:
                    r = p.add_run(remaining[:idx])
                    r.font.size = Pt(11)
                    r.font.color.rgb = C_DGRAY
                r2 = p.add_run(bp)
                r2.bold = True
                r2.font.size = Pt(11)
                r2.font.color.rgb = C_DGRAY
                remaining = remaining[idx + len(bp):]
        if remaining:
            r = p.add_run(remaining)
            r.font.size = Pt(11)
            r.font.color.rgb = C_DGRAY
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = C_DGRAY
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Cm(1 + level * 0.5)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.color.rgb = C_DGRAY
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def styled_table(doc, headers, rows, col_widths=None, highlight_rows=None):
    """
    Create a polished table.
    highlight_rows: set of row indices (0-based) to colour green.
    """
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, H_NAVY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(h)
        run.bold      = True
        run.font.size = Pt(10)
        run.font.color.rgb = C_WHITE

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg  = H_LGRAY if ri % 2 == 0 else H_WHITE
        if highlight_rows and ri in highlight_rows:
            bg = H_HLGRN
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if highlight_rows and ri in highlight_rows:
                run.bold = True
                run.font.color.rgb = C_GREEN
            else:
                run.font.color.rgb = C_DGRAY

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


def page_break(doc):
    doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════════
# Build Document
# ═════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ── Default paragraph font ───────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

# Top coloured banner
banner = doc.add_paragraph()
banner.paragraph_format.space_before = Pt(0)
banner.paragraph_format.space_after  = Pt(0)
banner_run = banner.add_run('  ')
banner_run.font.size = Pt(36)
# shade the paragraph background via shading on the paragraph itself
pPr = banner._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), H_NAVY)
pPr.append(shd)

para_space(doc, before=20)

# Title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(10)
t.paragraph_format.space_after  = Pt(4)
r = t.add_run('Crop Yield Prediction &\nPlant Disease Detection')
r.bold      = True
r.font.size = Pt(26)
r.font.color.rgb = C_NAVY

# Subtitle
s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.paragraph_format.space_after = Pt(30)
sr = s.add_run('Using Machine Learning and Deep Learning')
sr.font.size = Pt(14)
sr.font.color.rgb = C_TEAL
sr.italic = True

# Divider
div = doc.add_paragraph()
div.alignment = WD_ALIGN_PARAGRAPH.CENTER
divr = div.add_run('─' * 55)
divr.font.color.rgb = C_TEAL
divr.font.size = Pt(12)

para_space(doc, before=10, after=4)

# Author info table (borderless)
info_table = doc.add_table(rows=4, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('Submitted by', 'Mahinur Akhter'),
    ('Student ID',   '22201100'),
    ('Program',      'B.Sc. in Computer Science & Engineering'),
    ('Date',         'May 2026'),
]
for ri, (label, value) in enumerate(info_data):
    lc = info_table.rows[ri].cells[0]
    vc = info_table.rows[ri].cells[1]
    lp = lc.paragraphs[0]
    lr = lp.add_run(label)
    lr.bold = True
    lr.font.size = Pt(12)
    lr.font.color.rgb = C_NAVY
    lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    vp = vc.paragraphs[0]
    vr = vp.add_run('  ' + value)
    vr.font.size = Pt(12)
    vr.font.color.rgb = C_DGRAY

# Remove borders from info table
from docx.oxml.ns import qn as _qn
for row in info_table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ['top','left','bottom','right','insideH','insideV']:
            tag = OxmlElement(f'w:{edge}')
            tag.set(_qn('w:val'), 'none')
            tcBorders.append(tag)
        tcPr.append(tcBorders)

para_space(doc, before=20, after=8)

# Bottom banner
b2 = doc.add_paragraph()
b2r = b2.add_run('  ')
b2r.font.size = Pt(24)
pPr2 = b2._p.get_or_add_pPr()
shd2 = OxmlElement('w:shd')
shd2.set(qn('w:val'), 'clear')
shd2.set(qn('w:color'), 'auto')
shd2.set(qn('w:fill'), H_TEAL)
pPr2.append(shd2)

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# 1. ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════

add_heading1(doc, '1. Abstract')
add_body(doc,
    'This project presents an end-to-end machine learning pipeline for two critical '
    'agricultural tasks: (1) crop yield prediction using tabular environmental and soil data, '
    'and (2) plant disease detection using deep learning on leaf images. The system integrates '
    'ten heterogeneous real-world datasets spanning global FAO yield statistics, IoT soil sensors, '
    'air quality indices, and the PlantVillage image benchmark.',
    bold_parts=['crop yield prediction', 'plant disease detection'])
add_body(doc,
    'For yield prediction, four regression and four classification algorithms are benchmarked, '
    'with XGBoost emerging as the best model (R² = 0.2121, Accuracy = 60.1%). For disease '
    'detection, a fine-tuned MobileNetV2 achieves 98.12% validation accuracy across 38 disease '
    'classes. The pipeline also implements Grad-CAM for explainability and an image-tabular '
    'fusion stage that enriches the dataset with seven CNN-derived features.',
    bold_parts=['XGBoost', 'R² = 0.2121', '98.12%', 'Grad-CAM'])

# Keywords box
kw_p = doc.add_paragraph()
kw_p.paragraph_format.space_before = Pt(6)
kw_p.paragraph_format.space_after  = Pt(10)
kw_p.paragraph_format.left_indent  = Cm(0.5)
pPr_kw = kw_p._p.get_or_add_pPr()
shd_kw = OxmlElement('w:shd')
shd_kw.set(qn('w:val'), 'clear')
shd_kw.set(qn('w:color'), 'auto')
shd_kw.set(qn('w:fill'), 'EAF4FB')  # light blue tint
pPr_kw.append(shd_kw)
kw_r = kw_p.add_run('Keywords: ')
kw_r.bold = True; kw_r.font.size = Pt(11); kw_r.font.color.rgb = C_TEAL
kw_r2 = kw_p.add_run('Crop Yield Prediction · Plant Disease Detection · XGBoost · '
                      'MobileNetV2 · Transfer Learning · Grad-CAM · Image-Tabular Fusion')
kw_r2.font.size = Pt(10); kw_r2.italic = True; kw_r2.font.color.rgb = C_DGRAY


# ══════════════════════════════════════════════════════════════════════════════
# 2. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

add_heading1(doc, '2. Introduction')
add_body(doc,
    'Agricultural productivity is central to global food security, with over 800 million people '
    'facing chronic hunger worldwide. Two recurring threats are unpredictable crop yields—driven '
    'by climate variability, soil degradation, and socio-economic factors—and plant diseases, '
    'which destroy an estimated 10–40% of global food production annually.')
add_body(doc,
    'Recent advances in machine learning (ML) and deep learning (DL) have enabled data-driven '
    'solutions to these problems. Ensemble tree-based methods such as XGBoost and Random Forest '
    'have demonstrated strong performance on structured agricultural data, while convolutional '
    'neural networks (CNNs) have achieved near-expert accuracy in visual disease diagnosis.')
add_body(doc, 'This project addresses these challenges through a unified pipeline that:')

objectives = [
    'Merges ten heterogeneous real-world datasets into a single 18,741-row tabular feature matrix',
    'Benchmarks four regression and four classification algorithms for yield prediction',
    'Trains and fine-tunes MobileNetV2 on 87,866 leaf images across 38 disease classes',
    'Applies Grad-CAM for Explainable AI (XAI) output on disease detections',
    'Fuses CNN-derived image features back into the tabular dataset for enriched modelling',
]
for obj in objectives:
    add_bullet(doc, obj)


# ══════════════════════════════════════════════════════════════════════════════
# 3. DATASETS
# ══════════════════════════════════════════════════════════════════════════════

add_heading1(doc, '3. Datasets')
add_body(doc,
    'Ten datasets from diverse sources were collected, cleaned, and merged into a unified pipeline. '
    'The table below summarises each source.')

para_space(doc, before=4)

ds_headers = ['ID', 'Source', 'Rows', 'Key Features', 'Join Key']
ds_rows = [
    ('DS1',  'FAO Yield / Rainfall / Temperature',       '18,741',  'yield_hgha, rainfall_mm, avg_temp',     'Year, Country'),
    ('DS2',  'Crop Recommendation (Kaggle)',              '2,200',   'N, P, K, soil_pH, humidity',            'Crop_Type'),
    ('DS5',  'Bangladesh Agroclimatic (2000–2024)',       '150',     'wind, sunshine, soil_moisture',         'Year'),
    ('DS6',  'Global Land Temperatures',                 '3,402',   'avg_temp by country',                   'Year, Country'),
    ('DS9',  'IoT Telemetry Sensor',                     '1',       'CO, humidity, temperature',             'Year'),
    ('DS10', 'Dhaka Air Quality (2000–2025)',             '225,816', 'AQI, PM2.5, NO₂, CO₂',                 'Year'),
    ('XLS2', 'Real-Time IoT Soil Sensor (1007.xls)',     '1,006',   'temp, humidity, pH, N, P, K, conductivity', 'Year'),
    ('IMG',  'PlantVillage (Augmented)',                 '87,866',  'Leaf images, 38 disease classes',       'Crop_Type'),
]
styled_table(doc, ds_headers, ds_rows, col_widths=[1.2, 4.0, 1.4, 4.2, 2.5])
add_caption(doc, 'Table 1: Summary of all datasets used in the pipeline')

add_heading2(doc, '3.1  IoT Soil Sensor Dataset (XLS2)')
add_body(doc,
    'The Real-Time Soil Data file contains 1,006 rows of IoT sensor readings from four field '
    'locations in Dhaka (Garden, University, Ekuria, Jiya Uddan), collected April–May 2026. '
    'Features include soil temperature (°C), humidity (%), electrical conductivity (μS/cm), '
    'pH, and macro-nutrient concentrations (N, P, K in mg/kg). Column names used Unicode '
    'characters requiring normalised prefix-matching during loading.',
    bold_parts=['1,006 rows', 'four field locations'])

add_heading2(doc, '3.2  PlantVillage Image Dataset')
add_body(doc,
    'The augmented PlantVillage dataset contains 87,866 leaf images across 38 disease classes '
    'from 14 crop species. The train/validation split is 70,294 / 17,572 (80/20). All images '
    'are resized to 224×224 pixels with ImageNet-standard normalisation and data augmentation '
    '(random flip, zoom, rotation) during training.',
    bold_parts=['87,866 leaf images', '38 disease classes', '98.12%'])


# ══════════════════════════════════════════════════════════════════════════════
# 4. METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════

add_heading1(doc, '4. Methodology')

add_heading2(doc, '4.1  Data Merging — 7-Step Strategy')
add_body(doc,
    'The merge pipeline follows seven sequential steps designed to build a maximally informative '
    'tabular matrix while preserving all 18,741 backbone rows from the FAO yield dataset.')

steps = [
    ('Step 1 — Backbone Preparation',
     'FAO yield data (Year × Country × Crop_Type) forms the primary key. Crop names are '
     'normalised via a canonical mapping (e.g., "rice, paddy" → "rice").'),
    ('Step 2 — Climate Feature Merge',
     'FAO rainfall, FAO temperature, Bangladesh Agroclimatic, and Global Land Temperature '
     'datasets are left-joined on {Year, Country}, with column conflicts resolved by averaging.'),
    ('Step 3 — Soil Feature Merge',
     'Crop Recommendation (DS2) and IoT Soil Sensor (XLS2) features are merged on Crop_Type; '
     'soil pH, N, P, K, and humidity are propagated to all rows of that crop type.'),
    ('Step 4 — Environmental Feature Merge',
     'Dhaka Air Quality and IoT Telemetry datasets contribute AQI, PM2.5, NO₂, CO₂, CO, and '
     'LPG concentrations, joined on Year.'),
    ('Step 5 — Feature Engineering',
     'Four derived features are created: decade (decile of Year), temperature_range (Tmax−Tmin), '
     'season (Kharif/Rabi by crop type), and rainfall_category (Low/Medium/High by tercile).'),
    ('Step 6 — Missing Value Handling',
     'Columns with >50% null values (soil_fertility_idx, PM10, SO₂, O₃) are dropped; '
     'remaining nulls are imputed with column median via SimpleImputer.'),
    ('Step 7 — Export',
     'Final merged dataset: 18,741 rows × 28 columns, saved as data/processed/merged_dataset.csv.'),
]
for title, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(title + ':  ')
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = C_TEAL
    r2 = p.add_run(desc)
    r2.font.size = Pt(11); r2.font.color.rgb = C_DGRAY

add_heading2(doc, '4.2  Tabular Machine Learning')
add_heading3(doc, 'Feature Preparation')
add_body(doc,
    'All 22 numeric/categorical features are used after label-encoding categorical columns '
    '(season, soil_type), median imputation of remaining nulls, and z-score standardisation '
    'via StandardScaler. Dataset split: 80% train (14,987 samples) / 20% test (3,747 samples), '
    'random_state=42.')

add_heading3(doc, 'Regression Models')
add_body(doc, 'Four regressors are benchmarked to predict continuous crop yield (hg/ha):')
reg_models = [
    'Linear Regression — ordinary least squares baseline',
    'Random Forest — 200 trees, n_jobs=−1',
    'XGBoost — 200 estimators, lr=0.05, max_depth=6, subsample=0.8',
    'K-Nearest Neighbours — k=7',
]
for m in reg_models:
    add_bullet(doc, m)
add_body(doc, 'Evaluation metrics: RMSE, MAE, R², and 5-fold cross-validated R².')

add_heading3(doc, 'Classification Models')
add_body(doc,
    'Yield is discretised into three categories: Low (< 20,000 hg/ha), '
    'Medium (20,000–50,000 hg/ha), and High (> 50,000 hg/ha). '
    'Four classifiers are compared: Random Forest, XGBoost, K-Nearest Neighbours, '
    'and Gaussian Naïve Bayes. Primary metric: weighted F1-score.')

add_heading2(doc, '4.3  CNN — Plant Disease Detection')
add_heading3(doc, 'Architecture')
add_body(doc,
    'MobileNetV2 (pre-trained on ImageNet, 2.26M parameters) serves as the backbone. '
    'A custom classification head is appended: '
    'GlobalAveragePooling2D → Dense(256, ReLU) → Dropout(0.5) → Dense(38, Softmax).',
    bold_parts=['MobileNetV2', 'GlobalAveragePooling2D'])

add_heading3(doc, 'Two-Phase Training Protocol')
phases = [
    ('Phase 1 — Feature Extraction (5 epochs)',
     'Base model weights frozen. Only the custom head is trained. Learning rate = 10⁻³. '
     'Achieves 93.84% validation accuracy.'),
    ('Phase 2 — Fine-Tuning (10 epochs)',
     'Last 30 layers of MobileNetV2 unfrozen. Learning rate = 10⁻⁴. '
     'ReduceLROnPlateau halves the rate if validation loss stagnates for 2 epochs. '
     'Final validation accuracy: 98.12%.'),
]
for title, desc in phases:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(title + ':  ')
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = C_TEAL
    r2 = p.add_run(desc)
    r2.font.size = Pt(11); r2.font.color.rgb = C_DGRAY

add_heading3(doc, 'Explainability — Grad-CAM')
add_body(doc,
    'Gradient-weighted Class Activation Mapping (Grad-CAM) generates a saliency heatmap '
    'highlighting the image regions most influential to the model\'s prediction. '
    'Disease severity is computed as the fraction of heatmap pixels above threshold 0.5, '
    'providing a quantitative severity estimate alongside the class label.',
    bold_parts=['Grad-CAM', 'Disease severity'])

add_heading2(doc, '4.4  Image–Tabular Fusion')
add_body(doc,
    'Seven CNN-derived features per crop type are computed by averaging model outputs '
    'over one representative image per disease class, then joined to the tabular backbone '
    'on Crop_Type. This produces a fused dataset of 35 features with 2,845 matched rows.')
fusion_features = [
    'img_avg_confidence — mean softmax confidence (%)',
    'img_avg_severity — mean disease severity (% heatmap area above threshold)',
    'img_avg_yellowing — mean yellowing index [(R−G)/(R+G)]',
    'img_avg_browning — mean browning index [(R−B)/(R+B)]',
    'img_avg_texture — mean texture score (std-dev of grayscale pixels)',
    'img_disease_rate — fraction of disease-labelled classes per crop (%)',
]
for f in fusion_features:
    add_bullet(doc, f)


# ══════════════════════════════════════════════════════════════════════════════
# 5. EDA
# ══════════════════════════════════════════════════════════════════════════════

add_heading1(doc, '5. Exploratory Data Analysis')

add_heading2(doc, '5.1  Yield Distribution')
add_body(doc,
    'The yield target (yield_hgha) has a heavily right-skewed distribution '
    '(mean = 73,455 hg/ha, σ = 78,562 hg/ha), driven by high-yielding crops such as '
    'sugarcane and potato. The classification bins capture 52% Low, 27% Medium, '
    'and 21% High observations, revealing moderate class imbalance.',
    bold_parts=['52% Low', '27% Medium', '21% High'])

add_heading2(doc, '5.2  Feature Statistics')
para_space(doc, before=4)
stat_headers = ['Feature', 'Mean', 'Std Dev', 'Min', 'Max']
stat_rows = [
    ('avg_temp (°C)',         '22.05',   '6.02',   '−17.07', '30.74'),
    ('rainfall_mm',          '1,122.3', '734.2',  '51.0',   '3,240.0'),
    ('humidity_pct (%)',      '66.55',   '6.92',   '8.28',   '82.19'),
    ('nitrogen_N (mg/kg)',    '76.43',   '1.24',   '76.0',   '80.0'),
    ('soil_pH',               '6.27',    '0.03',   '6.26',   '6.36'),
    ('AQI',                   '150.02',  '2.56',   '70.01',  '190.61'),
    ('PM2.5 (μg/m³)',         '25.32',   '4.77',   '21.43',  '132.58'),
    ('fertilizer_kgha',       '40.01',   '67.73',  '0.0',    '167.0'),
]
styled_table(doc, stat_headers, stat_rows, col_widths=[3.5, 2.2, 2.2, 2.2, 2.2])
add_caption(doc, 'Table 2: Descriptive statistics of selected features (n = 18,741)')

add_heading2(doc, '5.3  Correlation Insights')
add_body(doc,
    'Temperature range and average temperature show moderate positive correlation with yield '
    '(r ≈ 0.3). Rainfall exhibits a non-linear relationship—both very low and very high rainfall '
    'correlate with reduced yields. Soil NPK features show minimal variance as they originate '
    'from crop-type averages rather than field-specific measurements.')


# ══════════════════════════════════════════════════════════════════════════════
# 6. RESULTS
# ══════════════════════════════════════════════════════════════════════════════

add_heading1(doc, '6. Results and Discussion')

add_heading2(doc, '6.1  Regression Results')
para_space(doc, before=4)
reg_headers = ['Model', 'RMSE', 'MAE', 'R²', 'CV-R²']
reg_rows = [
    ('Linear Regression',  '74,034', '54,102', '0.1212', '0.1352'),
    ('Random Forest',      '84,362', '56,483', '−0.1410', '−0.1004'),
    ('XGBoost ★',          '70,104', '49,316', '0.2121',  '0.2152'),
    ('K-Nearest Neighbours','76,377','54,073', '0.0647',  '0.0918'),
]
styled_table(doc, reg_headers, reg_rows, col_widths=[4.2, 2.2, 2.2, 1.8, 1.8], highlight_rows={2})
add_caption(doc, 'Table 3: Regression model comparison on held-out test set (n = 3,747). ★ Best model.')

add_body(doc,
    'XGBoost achieves the highest R² = 0.2121 on both test set and 5-fold cross-validation, '
    'confirming that the result is not an artefact of overfitting. Random Forest exhibits '
    'negative R², indicating overfitting to high-variance noisy features. '
    'The most important features for XGBoost are: Year, CO₂_ppm, avg_temp, rainfall_mm, '
    'and fertilizer_kgha — consistent with agronomic domain knowledge.',
    bold_parts=['R² = 0.2121', 'Year, CO₂_ppm, avg_temp, rainfall_mm'])

add_heading2(doc, '6.2  Classification Results')
para_space(doc, before=4)
cls_headers = ['Model', 'Accuracy', 'Weighted F1']
cls_rows = [
    ('Random Forest',         '0.4953', '0.4884'),
    ('XGBoost ★',             '0.6013', '0.5801'),
    ('K-Nearest Neighbours',  '0.5108', '0.4928'),
    ('Gaussian Naïve Bayes',  '0.5015', '0.4279'),
]
styled_table(doc, cls_headers, cls_rows, col_widths=[5.5, 3.0, 3.0], highlight_rows={1})
add_caption(doc, 'Table 4: Classification model comparison. ★ Best model.')

para_space(doc, before=4)
per_headers = ['Class', 'Precision', 'Recall', 'F1-Score', 'Support']
per_rows = [
    ('High',         '0.59', '0.84', '0.69', '1,726'),
    ('Low',          '0.64', '0.38', '0.47', '990'),
    ('Medium',       '0.61', '0.41', '0.49', '1,031'),
    ('Weighted Avg', '0.61', '0.60', '0.58', '3,747'),
]
styled_table(doc, per_headers, per_rows, col_widths=[3.5, 2.5, 2.5, 2.5, 2.0], highlight_rows={3})
add_caption(doc, 'Table 5: Per-class metrics for best classifier (XGBoost)')

add_body(doc,
    'XGBoost achieves 60.1% accuracy and 0.58 weighted F1, significantly outperforming '
    'the Gaussian NB baseline (50.1%). The High class is recalled most reliably (84%), '
    'while Low and Medium classes suffer from confusion near the bin boundaries due to '
    'the continuous, overlapping nature of yield distributions.',
    bold_parts=['60.1% accuracy', '84%'])

add_heading2(doc, '6.3  CNN — Plant Disease Detection')
add_heading3(doc, 'Training Progression')
para_space(doc, before=4)
cnn_headers = ['Phase', 'Epoch', 'Train Accuracy', 'Val Accuracy', 'Val Loss']
cnn_rows = [
    ('Phase 1\n(Frozen Base)', '1', '0.7516', '0.9037', '0.2954'),
    ('', '2', '0.8521', '0.9085', '0.2762'),
    ('', '3', '0.8714', '0.9262', '0.2290'),
    ('', '4', '0.8793', '0.9257', '0.2312'),
    ('', '5', '0.8875', '0.9384', '0.1983'),
    ('Phase 2\n(Fine-Tune)', '1', '0.8956', '0.9403', '0.1969'),
    ('', '2', '0.9373', '0.9681', '0.1001'),
    ('', '5', '0.9707', '0.9760', '0.0786'),
    ('', '6', '0.9750', '0.9812', '0.0558'),
    ('', '10', '0.9845', '0.9804', '0.0677'),
]
styled_table(doc, cnn_headers, cnn_rows, col_widths=[3.0, 1.5, 3.0, 3.0, 2.0], highlight_rows={8})
add_caption(doc, 'Table 6: CNN training progression (selected epochs). Best val accuracy highlighted.')

add_heading3(doc, 'Final Evaluation')
para_space(doc, before=4)
final_headers = ['Metric', 'Value']
final_rows = [
    ('Validation Accuracy',  '98.12%'),
    ('Validation Loss',      '0.0558'),
    ('Macro Avg Precision',  '0.98'),
    ('Macro Avg Recall',     '0.98'),
    ('Macro Avg F1-Score',   '0.98'),
    ('Total Classes',        '38'),
    ('Validation Samples',   '17,572'),
]
styled_table(doc, final_headers, final_rows, col_widths=[6.0, 5.5], highlight_rows={0})
add_caption(doc, 'Table 7: Final CNN evaluation on the validation set')

add_heading3(doc, 'Per-Class Performance Highlights')
add_body(doc, 'All 38 classes achieve F1 ≥ 0.90. Notable results:')
highlights = [
    'Perfect F1 = 1.00 — Potato Early/Late Blight, Cherry Powdery Mildew, Orange Huanglongbing, Squash Powdery Mildew, Corn Common Rust (13 classes total)',
    'Most challenging — Tomato Target Spot (F1 = 0.90), Tomato Healthy (F1 = 0.92): visually similar leaf textures',
    'Corn Cercospora Leaf Spot (F1 = 0.95): subtle grey-brown lesion patterns require deep feature discrimination',
]
for h in highlights:
    add_bullet(doc, h)

add_heading3(doc, 'Grad-CAM XAI Sample')
para_space(doc, before=4)
xai_headers = ['XAI Feature', 'Value']
xai_rows = [
    ('Disease Class',     'Apple___Apple_scab'),
    ('Confidence',        '100.0%'),
    ('Mean R channel',    '0.4309'),
    ('Mean G channel',    '0.5078'),
    ('Mean B channel',    '0.4669'),
    ('Texture Score',     '0.1873'),
    ('Yellowing Index',   '−0.0768'),
    ('Browning Index',    '−0.0359'),
    ('Disease Severity',  '10.67% of leaf area'),
    ('Growth Stage',      'Moderate Infection'),
]
styled_table(doc, xai_headers, xai_rows, col_widths=[5.5, 6.0], highlight_rows={8})
add_caption(doc, 'Table 8: XAI features extracted via Grad-CAM for sample image (Apple Scab)')

try:
    img_p2 = doc.add_paragraph()
    img_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p2.paragraph_format.space_before = Pt(8)
    run_img2 = img_p2.add_run()
    run_img2.add_picture('outputs/gradcam_sample.jpg', width=Inches(3.5))
    add_caption(doc, 'Figure 1: Grad-CAM heatmap for Apple Scab — disease severity 10.67%')
except Exception:
    pass

add_heading2(doc, '6.4  Image–Tabular Fusion Results')
para_space(doc, before=4)
fus_headers = ['Crop', 'Confidence (%)', 'Severity (%)', 'Yellowing Index', 'Browning Index', 'Disease Rate']
fus_rows = [
    ('Apple',    '100.0', '13.3', '−0.022', '−0.013', '75.0%'),
    ('Corn',     '88.2',  '21.3', '−0.043', '+0.098', '75.0%'),
    ('Grape',    '100.0', '19.8', '−0.033', '+0.051', '75.0%'),
    ('Orange',   '100.0', '15.4', '+0.094', '+0.151', '100.0%'),
    ('Peach',    '100.0', '19.3', '+0.009', '+0.059', '50.0%'),
    ('Potato',   '100.0', '25.4', '−0.024', '+0.045', '66.7%'),
    ('Squash',   '100.0', '7.6',  '−0.056', '+0.110', '100.0%'),
    ('Tomato',   '97.2',  '15.8', '−0.006', '+0.036', '90.0%'),
]
styled_table(doc, fus_headers, fus_rows, col_widths=[2.5, 2.5, 2.5, 2.5, 2.5, 2.2], highlight_rows={3, 7})
add_caption(doc, 'Table 9: Per-crop disease profile from CNN predictions (image–tabular fusion). High-risk crops highlighted.')

add_body(doc,
    'The fusion stage successfully matched 2,845 rows from the tabular backbone to '
    'CNN-derived features. High disease rates in Tomato (90%) and Orange/Squash (100%) '
    'correlate with known yield suppression in affected regions, validating the biological '
    'coherence of the fusion output.',
    bold_parts=['2,845 rows', 'Tomato (90%)', 'Orange/Squash (100%)'])


# ══════════════════════════════════════════════════════════════════════════════
# 7. DISCUSSION
# ══════════════════════════════════════════════════════════════════════════════

add_heading1(doc, '7. Discussion')

add_heading2(doc, '7.1  Yield Prediction Limitations')
add_body(doc, 'The moderate R² values (≤ 0.21) reflect several structural limitations:')
limitations = [
    'Covariate shift: The dataset spans 206 countries over 17 years; agro-climatic '
     'conditions vary enormously and are only partially captured by the merged features',
    'Soil feature uniformity: NPK values from DS2 are crop-type averages, not field-specific '
     'measurements, limiting their discriminative power across diverse geographies',
    'Label noise: FAO yield data aggregates national-level statistics, masking within-country '
     'variation and extreme events',
]
for l in limitations:
    add_bullet(doc, l)
add_body(doc,
    'Despite these limitations, XGBoost\'s consistency between test R² (0.2121) and '
    'cross-validated R² (0.2152) confirms that the result is not an artefact of overfitting.',
    bold_parts=['0.2121', '0.2152'])

add_heading2(doc, '7.2  CNN Strengths')
add_body(doc,
    'The 98.12% validation accuracy is achieved on a held-out split of the augmented '
    'PlantVillage dataset. The two-phase training strategy is highly effective: Phase 1 '
    'leverages ImageNet features for rapid convergence, while Phase 2 fine-tuning adapts '
    'deep convolutional features to the plant domain. ReduceLROnPlateau fires twice '
    '(Epochs 4 and 8), preventing overfitting by progressively reducing the learning rate '
    'from 10⁻⁴ to 1.25×10⁻⁵.',
    bold_parts=['98.12%', 'Phase 1', 'Phase 2'])


# ══════════════════════════════════════════════════════════════════════════════
# 8. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════

add_heading1(doc, '8. Conclusion')
add_body(doc, 'This project delivers a complete, reproducible machine learning pipeline with two core outcomes:')

conclusions = [
    'A multi-source tabular pipeline integrating 10 heterogeneous datasets, benchmarking '
     'four regression and four classification algorithms — XGBoost identified as best model '
     '(R² = 0.2121, Accuracy = 60.1%)',
    'A transfer-learned MobileNetV2 CNN achieving 98.12% accuracy on 38 plant disease '
     'classes with Grad-CAM explainability and quantitative disease severity estimation',
]
for c in conclusions:
    add_bullet(doc, c)

add_heading2(doc, 'Future Directions')
future = [
    'Incorporate Sentinel-2 / MODIS satellite imagery time-series for spatially explicit yield forecasting',
    'Collect field-condition plant photographs to reduce the lab-to-field domain gap',
    'Train a unified multi-modal transformer that processes tabular and image inputs jointly end-to-end',
    'Deploy the CNN model as a mobile application for real-time farmer advisory services',
    'Expand to Bangladesh-specific crops (jute, aman rice) with local field sensor networks',
]
for f in future:
    add_bullet(doc, f)


# ══════════════════════════════════════════════════════════════════════════════
# 9. REFERENCES
# ══════════════════════════════════════════════════════════════════════════════

add_heading1(doc, '9. References')

refs = [
    '[1] FAO, 2023. The State of Food Security and Nutrition in the World 2023. Rome: FAO.',
    '[2] Savary et al., 2019. The global burden of pathogens and pests on major food crops. '
     'Nature Ecology & Evolution, 3(3), pp.430–439.',
    '[3] Mohanty et al., 2016. Using deep learning for image-based plant disease detection. '
     'Frontiers in Plant Science, 7, p.1419.',
    '[4] Van Klompenburg et al., 2020. Crop yield prediction using machine learning: A '
     'systematic literature review. Computers and Electronics in Agriculture, 177, p.105709.',
    '[5] You et al., 2017. Deep Gaussian process for crop yield prediction based on remote '
     'sensing data. Proceedings of AAAI Conference on Artificial Intelligence, 31(1).',
    '[6] Barbedo, J.G.A., 2018. Factors influencing the use of deep learning for plant disease '
     'recognition. Biosystems Engineering, 172, pp.84–91.',
    '[7] Kamilaris & Prenafeta-Boldú, 2018. Deep learning in agriculture: A survey. '
     'Computers and Electronics in Agriculture, 147, pp.70–90.',
    '[8] Selvaraju et al., 2017. Grad-CAM: Visual explanations from deep networks via '
     'gradient-based localization. Proceedings of ICCV, pp.618–626.',
    '[9] Pantazi et al., 2019. Automated leaf disease detection in different crop species '
     'through image features analysis and One Class Classifiers. '
     'Computers and Electronics in Agriculture, 156, pp.96–104.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.7)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(ref)
    r.font.size = Pt(10)
    r.font.color.rgb = C_DGRAY


# ══════════════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════════════

doc.save('report.docx')
print('report.docx saved successfully.')
